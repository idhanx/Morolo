"""Document text extraction service with OCR support."""

import hashlib
import io
import logging
from dataclasses import dataclass
from typing import BinaryIO

import pytesseract
from docx import Document as DocxDocument
from pdf2image import convert_from_bytes
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdfminer.pdfpage import PDFPage
from PIL import Image, ImageEnhance, ImageFilter
from redis import Redis

from backend.core.config import settings
from backend.core.types import ScanType

logger = logging.getLogger(__name__)

# All major Indian languages + English for Tesseract
INDIAN_OCR_LANGS = (
    "eng+hin+tam+tel+kan+mal+mar+guj+pan+ben+ori+urd"
)


@dataclass
class ExtractedText:
    """Result of text extraction from a document."""
    text: str
    scan_type: ScanType
    page_count: int
    char_count: int
    cached: bool = False


def _preprocess_image(image: Image.Image) -> Image.Image:
    """Preprocess image to improve OCR accuracy on government IDs.
    
    Steps:
    1. Convert to grayscale
    2. Increase contrast (helps with faded/low-quality scans)
    3. Sharpen (helps with blurry images)
    4. Scale up if small (Tesseract works better on larger images)
    """
    # Convert to grayscale
    image = image.convert("L")
    
    # Increase contrast
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(2.0)
    
    # Sharpen
    image = image.filter(ImageFilter.SHARPEN)
    
    # Scale up if image is small — Tesseract needs at least 300 DPI equivalent
    width, height = image.size
    if width < 1000 or height < 1000:
        scale = max(1000 / width, 1000 / height)
        new_size = (int(width * scale), int(height * scale))
        image = image.resize(new_size, Image.LANCZOS)
        logger.info(f"Scaled image from {width}x{height} to {new_size[0]}x{new_size[1]}")
    
    return image


def _ocr_image(image: Image.Image) -> str:
    """Run Tesseract OCR with Indian language support and fallback."""
    preprocessed = _preprocess_image(image)
    
    # Try with all Indian languages first
    try:
        text = pytesseract.image_to_string(
            preprocessed,
            lang=INDIAN_OCR_LANGS,
            config="--oem 3 --psm 3",  # OEM 3=LSTM, PSM 3=auto page segmentation
        )
        if text.strip():
            logger.info(f"OCR with Indian langs: {len(text)} chars")
            return text
    except Exception as e:
        logger.warning(f"Indian language OCR failed: {e}, falling back to English")
    
    # Fallback to English only
    try:
        text = pytesseract.image_to_string(
            preprocessed,
            lang="eng",
            config="--oem 3 --psm 3",
        )
        if text.strip():
            logger.info(f"OCR with English fallback: {len(text)} chars")
            return text
    except Exception as e:
        logger.error(f"OCR English fallback failed: {e}")
    
    # All OCR attempts failed
    logger.error("OCR completely failed for image. No text extracted.")
    raise ValueError("Unable to extract text from image. OCR processing failed. File may be corrupted or unreadable.")


class DocumentProcessor:
    """Service for extracting text from various document formats."""

    def __init__(self, redis_client: Redis | None = None):
        self.redis_client = redis_client
        self.ocr_cache_ttl = 3600
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    # ------------------------------------------------------------------
    # Caching helpers
    # ------------------------------------------------------------------

    def _compute_file_hash(self, file_bytes: bytes) -> str:
        return hashlib.sha256(file_bytes).hexdigest()

    def _get_cached_text(self, file_hash: str) -> str | None:
        if not self.redis_client:
            return None
        try:
            cached = self.redis_client.get(f"ocr:{file_hash}")
            if cached:
                logger.info(f"Cache hit for file hash {file_hash[:8]}...")
                return cached.decode("utf-8")
        except Exception as e:
            logger.warning(f"Failed to retrieve from cache: {e}")
        return None

    def _cache_text(self, file_hash: str, text: str) -> None:
        if not self.redis_client:
            return
        try:
            self.redis_client.setex(
                f"ocr:{file_hash}", self.ocr_cache_ttl, text.encode("utf-8")
            )
            logger.info(f"Cached OCR result for {file_hash[:8]}...")
        except Exception as e:
            logger.warning(f"Failed to cache OCR result: {e}")

    # ------------------------------------------------------------------
    # Scan type detection
    # ------------------------------------------------------------------

    def detect_scan_type(self, pdf_bytes: bytes) -> ScanType:
        """Detect if a PDF is text-based, scanned, or mixed."""
        try:
            pdf_stream = io.BytesIO(pdf_bytes)
            pages = list(PDFPage.get_pages(pdf_stream))
            if not pages:
                return ScanType.SCANNED

            threshold = settings.OCR_CHAR_THRESHOLD
            page_char_counts: list[int] = []

            for i in range(len(pages)):
                page_stream = io.BytesIO(pdf_bytes)
                text = pdfminer_extract_text(page_stream, page_numbers=[i])
                page_char_counts.append(len(text.strip()))

            avg_chars = sum(page_char_counts) / len(page_char_counts)

            if avg_chars >= threshold:
                return ScanType.TEXT
            if all(c < threshold for c in page_char_counts):
                return ScanType.SCANNED
            return ScanType.MIXED

        except Exception as e:
            logger.error(f"Failed to detect scan type: {e}")
            return ScanType.SCANNED

    # ------------------------------------------------------------------
    # PDF extraction
    # ------------------------------------------------------------------

    def _extract_text_from_pdf(self, pdf_bytes: bytes) -> tuple[str, ScanType]:
        scan_type = self.detect_scan_type(pdf_bytes)

        if scan_type == ScanType.TEXT:
            try:
                text = pdfminer_extract_text(io.BytesIO(pdf_bytes))
                if text.strip():
                    logger.info("Extracted text from text-based PDF via pdfminer")
                    return text, scan_type
                # Empty text despite TEXT classification — force OCR
                logger.warning("pdfminer returned empty text, forcing OCR")
                scan_type = ScanType.SCANNED
            except Exception as e:
                logger.error(f"pdfminer extraction failed: {e}, falling back to OCR")
                scan_type = ScanType.SCANNED

        # Check cache before OCR
        file_hash = self._compute_file_hash(pdf_bytes)
        cached_text = self._get_cached_text(file_hash)
        if cached_text:
            return cached_text, scan_type

        # Convert PDF pages to images and OCR each one
        try:
            # Use high DPI for better OCR quality (300 DPI recommended)
            images = convert_from_bytes(pdf_bytes, dpi=300)
            logger.info(f"Converted PDF to {len(images)} images for OCR at 300 DPI")
            
            text_parts = []
            for i, image in enumerate(images, start=1):
                page_text = _ocr_image(image)
                text_parts.append(page_text)
                logger.debug(f"OCR page {i}/{len(images)}: {len(page_text)} chars")
            
            text = "\n\n".join(text_parts)
            # Normalize encoding artifacts from OCR
            text = text.encode("utf-8", errors="replace").decode("utf-8")
            self._cache_text(file_hash, text)
            logger.info(f"OCR complete: {len(text)} chars across {len(images)} pages")
            return text, scan_type

        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            raise RuntimeError(f"Failed to extract text from PDF: {e}")

    # ------------------------------------------------------------------
    # Image extraction
    # ------------------------------------------------------------------

    def _extract_text_from_image(self, image_bytes: bytes) -> str:
        file_hash = self._compute_file_hash(image_bytes)
        cached = self._get_cached_text(file_hash)
        if cached:
            return cached
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = _ocr_image(image)
            self._cache_text(file_hash, text)
            logger.info(f"Image OCR complete: {len(text)} chars")
            return text
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            raise RuntimeError(f"Failed to extract text from image: {e}")

    # ------------------------------------------------------------------
    # DOCX extraction
    # ------------------------------------------------------------------

    def _extract_text_from_docx(self, docx_bytes: bytes) -> str:
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            text = "\n".join(p.text for p in doc.paragraphs)
            logger.info(f"Extracted DOCX text: {len(text)} chars")
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise RuntimeError(f"Failed to extract text from DOCX: {e}")

    @staticmethod
    def _docx_page_count(docx_bytes: bytes) -> int:
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            total_chars = sum(len(p.text) for p in doc.paragraphs)
            return max(1, total_chars // 3000)
        except Exception:
            return 1

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def extract_text(self, file_bytes: bytes, filename: str) -> ExtractedText:
        """Extract text from a document based on its file extension."""
        file_extension = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if file_extension == "pdf":
            text, scan_type = self._extract_text_from_pdf(file_bytes)
            page_count = sum(1 for _ in PDFPage.get_pages(io.BytesIO(file_bytes)))

        elif file_extension in ("png", "jpg", "jpeg"):
            text = self._extract_text_from_image(file_bytes)
            scan_type = ScanType.SCANNED
            page_count = 1

        elif file_extension == "docx":
            text = self._extract_text_from_docx(file_bytes)
            scan_type = ScanType.TEXT
            page_count = self._docx_page_count(file_bytes)

        else:
            raise ValueError(f"Unsupported file type: {file_extension!r}")

        char_count = len(text)
        file_hash = self._compute_file_hash(file_bytes)
        cached = self._get_cached_text(file_hash) is not None

        return ExtractedText(
            text=text,
            scan_type=scan_type,
            page_count=page_count,
            char_count=char_count,
            cached=cached,
        )